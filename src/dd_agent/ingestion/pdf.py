"""Document text extraction — PDF, DOCX, plain text. Top-level dispatcher.

History: this module was named `pdf.py` and was PDF-only. It now also handles
DOCX (Microsoft Word) and plain-text files. The `extract_text()` PDF-only
function is kept for backward compat with callers that already know the file
is a PDF. New code should call `extract_document()` and let it dispatch by
file extension.

The v7 motivation: a Rivian memo arrived as `.docx`, the orchestrator routed
it through `Path.read_text()` (since the extension wasn't `.pdf`), and the
ZIP magic header `PK\\x03\\x04` got extracted as the literal company name "PK".
The dispatcher prevents that whole class of failures.
"""

from __future__ import annotations

import logging
from pathlib import Path

from pypdf import PdfReader

log = logging.getLogger("dd_agent.documents")


def extract_text(pdf_path: str | Path, max_pages: int | None = None) -> str:
    """Extract text from a PDF. Returns one string with form-feed page separators."""
    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF not found: {path}")
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        if max_pages is not None and i >= max_pages:
            break
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            pages.append(f"[page {i+1} extraction failed: {exc}]")
    return "\f".join(pages)


def page_count(pdf_path: str | Path) -> int:
    return len(PdfReader(str(pdf_path)).pages)


def extract_docx(docx_path: str | Path) -> str:
    """Extract text from a Microsoft Word .docx file.

    Iterates paragraphs and tables; concatenates with newline separators.
    Tables get cell text joined by tab + paragraph by newline. Returns
    a single string ready for downstream LLM normalization."""
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise RuntimeError(
            "python-docx not installed — run `pip install python-docx`"
        ) from exc
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")
    doc = docx.Document(str(path))
    blocks: list[str] = []
    # Paragraphs
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            blocks.append(text)
    # Tables (cells separated by tab)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            row_text = "\t".join(c for c in cells if c)
            if row_text:
                blocks.append(row_text)
    return "\n\n".join(blocks)


def extract_document(path: str | Path) -> str:
    """Top-level dispatcher. Picks the right extractor by extension.

    Supported:
      .pdf       → pypdf (existing extract_text)
      .docx      → python-docx (extract_docx)
      .md / .markdown / .txt → read_text()
      .doc       → raises NotImplementedError (legacy Word format, rarely seen)

    Unknown extensions fall back to read_text(errors='ignore') with a warning.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"file not found: {p}")
    ext = p.suffix.lower()
    if ext == ".pdf":
        return extract_text(p)
    if ext == ".docx":
        return extract_docx(p)
    if ext in (".md", ".markdown", ".txt"):
        return p.read_text(encoding="utf-8", errors="ignore")
    if ext == ".doc":
        raise NotImplementedError(
            f"Legacy .doc format not supported. Convert {p.name} to .docx or .pdf first."
        )
    log.warning("unknown file extension %r for %s — reading as text", ext, p)
    return p.read_text(encoding="utf-8", errors="ignore")
